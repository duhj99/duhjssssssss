import os
import re
from docx import Document
import zipfile
import tempfile
import shutil

class WordProcessor:
    @staticmethod
    def find_replace(file_path, find_text, replace_text, use_regex=False, output_dir=None):
        """
        在Word文档中查找并替换文本
        
        Args:
            file_path: Word文档路径
            find_text: 要查找的文本
            replace_text: 替换的文本
            use_regex: 是否使用正则表达式
            output_dir: 输出目录，如果为None则覆盖原文件
            
        Returns:
            处理后的文件路径
        """
        try:
            # 创建输出文件路径
            if output_dir:
                output_path = os.path.join(output_dir, os.path.basename(file_path))
            else:
                output_path = file_path
                
            # 如果输出路径与输入路径不同，先复制文件
            if output_path != file_path:
                shutil.copy2(file_path, output_path)
            
            # 打开Word文档
            doc = Document(output_path)
            
            # 替换文本
            for paragraph in doc.paragraphs:
                if use_regex:
                    paragraph.text = re.sub(find_text, replace_text, paragraph.text)
                else:
                    paragraph.text = paragraph.text.replace(find_text, replace_text)
            
            # 替换表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if use_regex:
                                paragraph.text = re.sub(find_text, replace_text, paragraph.text)
                            else:
                                paragraph.text = paragraph.text.replace(find_text, replace_text)
            
            # 保存文档
            doc.save(output_path)
            return output_path
        except Exception as e:
            raise Exception(f"处理Word文档时出错: {str(e)}")
    
    @staticmethod
    def batch_find_replace(file_paths, replacements, use_regex=False, output_dir=None):
        """
        批量处理多个Word文档的查找替换
        
        Args:
            file_paths: Word文档路径列表
            replacements: 替换规则列表，每个规则是一个(find_text, replace_text)元组
            use_regex: 是否使用正则表达式
            output_dir: 输出目录
            
        Returns:
            处理后的文件路径列表
        """
        processed_files = []
        
        for file_path in file_paths:
            # 为每个文件创建一个副本
            if output_dir:
                output_path = os.path.join(output_dir, os.path.basename(file_path))
                shutil.copy2(file_path, output_path)
            else:
                output_path = file_path
            
            # 应用所有替换规则
            doc = Document(output_path)
            
            for find_text, replace_text in replacements:
                # 替换段落中的文本
                for paragraph in doc.paragraphs:
                    if use_regex:
                        paragraph.text = re.sub(find_text, replace_text, paragraph.text)
                    else:
                        paragraph.text = paragraph.text.replace(find_text, replace_text)
                
                # 替换表格中的文本
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if use_regex:
                                    paragraph.text = re.sub(find_text, replace_text, paragraph.text)
                                else:
                                    paragraph.text = paragraph.text.replace(find_text, replace_text)
            
            # 保存文档
            doc.save(output_path)
            processed_files.append(output_path)
        
        return processed_files
    
    @staticmethod
    def merge_documents(file_paths, output_path):
        """
        合并多个Word文档
        
        Args:
            file_paths: Word文档路径列表
            output_path: 输出文件路径
            
        Returns:
            合并后的文件路径
        """
        try:
            # 创建一个新文档或使用第一个文档作为基础
            if file_paths:
                merged_doc = Document(file_paths[0])
                
                # 从第二个文档开始合并
                for i in range(1, len(file_paths)):
                    # 添加分页符
                    merged_doc.add_page_break()
                    
                    # 打开当前文档
                    doc = Document(file_paths[i])
                    
                    # 复制段落
                    for paragraph in doc.paragraphs:
                        merged_para = merged_doc.add_paragraph()
                        merged_para.text = paragraph.text
                        merged_para.style = paragraph.style
                    
                    # 复制表格
                    for table in doc.tables:
                        # 创建新表格
                        rows = len(table.rows)
                        cols = len(table.rows[0].cells) if rows > 0 else 0
                        merged_table = merged_doc.add_table(rows=rows, cols=cols)
                        
                        # 复制单元格内容
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                merged_table.rows[i].cells[j].text = cell.text
                
                # 保存合并后的文档
                merged_doc.save(output_path)
                return output_path
            else:
                raise Exception("没有提供要合并的文档")
        except Exception as e:
            raise Exception(f"合并Word文档时出错: {str(e)}")
    
    @staticmethod
    def extract_content(file_path, output_path, extract_type="text"):
        """
        从Word文档中提取内容
        
        Args:
            file_path: Word文档路径
            output_path: 输出文件路径
            extract_type: 提取类型，可以是"text"、"tables"或"all"
            
        Returns:
            提取的内容文件路径
        """
        try:
            doc = Document(file_path)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                # 提取文本
                if extract_type in ["text", "all"]:
                    f.write("# 文档文本内容\n\n")
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            f.write(paragraph.text + "\n")
                    f.write("\n\n")
                
                # 提取表格
                if extract_type in ["tables", "all"]:
                    f.write("# 文档表格内容\n\n")
                    for i, table in enumerate(doc.tables):
                        f.write(f"## 表格 {i+1}\n\n")
                        for row in table.rows:
                            f.write("| " + " | ".join(cell.text for cell in row.cells) + " |\n")
                        f.write("\n\n")
            
            return output_path
        except Exception as e:
            raise Exception(f"提取Word文档内容时出错: {str(e)}")